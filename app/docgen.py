import os
from typing import Dict
from docx import Document


def generate_resume_docx(base_template_path: str, out_path: str, data: Dict) -> None:
    doc = Document(base_template_path)

    replacements = {
        "{{FULL_NAME}}": data.get("full_name", ""),
        "{{EMAIL}}": data.get("email", ""),
        "{{PHONE}}": data.get("phone", ""),
        "{{LOCATION}}": data.get("location", ""),
        "{{LINKEDIN}}": data.get("linkedin", ""),
        "{{SUMMARY}}": data.get("summary", ""),
        "{{SKILLS}}": ", ".join(data.get("skills", [])),
    }

    # Replace template placeholders
    for p in doc.paragraphs:
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    # Insert highlights near the top (ATS-safe)
    doc.add_paragraph("")  # spacing
    doc.add_paragraph("TARGETED HIGHLIGHTS")
    for b in data.get("tailored_highlights", []):
        doc.add_paragraph(b, style="List Bullet")

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)